from flask import Flask, render_template, request, send_file, redirect, url_for, flash
import requests
from bs4 import BeautifulSoup
import io
import os
import openai
import PyPDF2
from docx import Document
from xhtml2pdf import pisa
import re
from uuid import uuid4

app = Flask(__name__)
app.secret_key = 'supersecretkey'  # Used for flash messages

# Set your OpenAI API key (ensure it is set in your environment)
openai.api_key = os.getenv("OPENAI_API_KEY")

# In-memory cache for files (for demonstration only – not for production)
cache = {}

def extract_job_description(url):
    """Download and extract text from a job description webpage."""
    try:
        response = requests.get(url)
        if response.status_code != 200:
            return None, "Failed to retrieve job description."
        soup = BeautifulSoup(response.text, 'html.parser')
        # Use paragraphs as heuristic; fallback to all text if needed.
        paragraphs = soup.find_all('p')
        text = "\n".join([p.get_text() for p in paragraphs])
        if not text.strip():
            text = soup.get_text()
        return text, None
    except Exception as e:
        return None, str(e)

def extract_resume_text(file_stream):
    """Extract text from an uploaded PDF resume using PyPDF2."""
    try:
        pdf_reader = PyPDF2.PdfReader(file_stream)
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text
    except Exception as e:
        return None

def generate_resume(original_resume, job_description):
    """Use OpenAI GPT to rewrite the resume to match the job description."""
    prompt = f"""
You are an expert resume writer. Given the following original resume and a job description, rewrite and optimize the resume to match the job requirements while ensuring it accurately reflects the candidate's skills. Generate new bullet points and rewrite sections as needed for clarity and ATS optimization.

Job Description:
{job_description}

Original Resume:
{original_resume}

Please provide the revised resume in plain text format with clearly marked sections (e.g., Summary, Experience, Education, Skills, etc.).
    """
    try:
        response = openai.Completion.create(
            engine="text-davinci-003",
            prompt=prompt,
            max_tokens=1500,
            temperature=0.7,
            top_p=1,
            n=1
        )
        revised_resume = response.choices[0].text.strip()
        return revised_resume
    except Exception as e:
        print("OpenAI API error:", e)
        return None

def compute_match_score(job_description, resume_text):
    """Compute a simple match score based on keyword overlap."""
    def preprocess(text):
        text = text.lower()
        text = re.sub(r'[^\w\s]', '', text)
        return set(text.split())
    jd_words = preprocess(job_description)
    resume_words = preprocess(resume_text)
    common_words = jd_words.intersection(resume_words)
    if len(jd_words) == 0:
        return 0
    score = (len(common_words) / len(jd_words)) * 100
    return round(score, 2)

def generate_docx(resume_text, match_score):
    """Generate a DOCX file from the revised resume text."""
    document = Document()
    document.add_heading("Tailored Resume", level=0)
    document.add_paragraph(f"Match Score: {match_score}%")
    for line in resume_text.split('\n'):
        if line.strip():
            document.add_paragraph(line.strip())
    docx_io = io.BytesIO()
    document.save(docx_io)
    docx_io.seek(0)
    return docx_io

def generate_pdf(html_content):
    """Generate a PDF file from HTML content using xhtml2pdf."""
    pdf_io = io.BytesIO()
    pisa_status = pisa.CreatePDF(io.StringIO(html_content), dest=pdf_io)
    if pisa_status.err:
        return None
    pdf_io.seek(0)
    return pdf_io

def render_resume_html(resume_text, match_score):
    """Create an HTML representation of the resume with ATS-friendly styling."""
    html_template = f"""
<html>
<head>
    <style>
        body {{
            font-family: Arial, sans-serif;
            margin: 40px;
        }}
        h1 {{
            text-align: center;
        }}
        .match-score {{
            font-size: 14px;
            color: #555;
            text-align: center;
        }}
        .content {{
            margin-top: 30px;
            white-space: pre-wrap;
        }}
    </style>
</head>
<body>
    <h1>Tailored Resume</h1>
    <p class="match-score">Match Score: {match_score}%</p>
    <div class="content">
        {resume_text}
    </div>
</body>
</html>
    """
    return html_template

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        job_url = request.form.get("job_url")
        resume_file = request.files.get("resume_file")
        
        if not job_url or not resume_file:
            flash("Please provide both a job description URL and a resume PDF.")
            return redirect(request.url)
        
        # Extract the job description text.
        job_description, error = extract_job_description(job_url)
        if error or not job_description:
            flash("Failed to extract job description from the provided URL.")
            return redirect(request.url)
        
        # Extract text from the uploaded PDF resume.
        resume_text = extract_resume_text(resume_file)
        if not resume_text:
            flash("Failed to extract text from the uploaded resume.")
            return redirect(request.url)
        
        # Generate the revised resume using OpenAI.
        revised_resume = generate_resume(resume_text, job_description)
        if not revised_resume:
            flash("Failed to generate a revised resume.")
            return redirect(request.url)
        
        # Compute a match score based on keyword overlap.
        match_score = compute_match_score(job_description, revised_resume)
        
        # Render the resume in an HTML template.
        resume_html = render_resume_html(revised_resume, match_score)
        
        # Generate DOCX and PDF files.
        docx_file = generate_docx(revised_resume, match_score)
        pdf_file = generate_pdf(resume_html)
        if not pdf_file:
            flash("Failed to generate PDF file.")
            return redirect(request.url)
        
        # Store files in the in-memory cache using a unique ID.
        file_id = str(uuid4())
        cache[file_id] = {
            "docx": docx_file,
            "pdf": pdf_file,
            "resume_html": resume_html,
            "revised_resume": revised_resume,
            "match_score": match_score
        }
        return redirect(url_for("download", file_id=file_id))
    return render_template("index.html")

@app.route("/download/<file_id>")
def download(file_id):
    if file_id not in cache:
        return "File not found.", 404
    data = cache[file_id]
    return render_template("download.html", file_id=file_id, match_score=data["match_score"])

@app.route("/download/pdf/<file_id>")
def download_pdf(file_id):
    if file_id not in cache:
        return "File not found.", 404
    pdf_file = cache[file_id]["pdf"]
    return send_file(pdf_file, as_attachment=True,
                     download_name="Tailored_Resume.pdf",
                     mimetype="application/pdf")

@app.route("/download/docx/<file_id>")
def download_docx(file_id):
    if file_id not in cache:
        return "File not found.", 404
    docx_file = cache[file_id]["docx"]
    return send_file(docx_file, as_attachment=True,
                     download_name="Tailored_Resume.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == "__main__":
    app.run(debug=True)
